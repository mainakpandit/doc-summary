"""Format-specific document parsers producing a normalized ParsedDocument.

Every accepted file (see README "Accepted formats") is parsed into one flat
`text` string plus a list of `Segment` spans over that string, so downstream
chunking (`services/ingestion.py`, not yet built) has a single shape to work
from regardless of source format. Parsing is CPU/IO-bound sync work (pypdf,
pdfplumber, python-docx are all sync libraries), so `parse_file` runs it in
a thread rather than blocking the event loop.
"""

import asyncio
import csv
import json
import re
from pathlib import Path

import magic
import pdfplumber
from docx import Document as DocxDocument
from pydantic import BaseModel
from pypdf import PdfReader


class Segment(BaseModel):
    char_start: int
    char_end: int
    page: int | None = None


class ParsedDocument(BaseModel):
    mime_type: str
    text: str
    segments: list[Segment]


class UnsupportedFormat(Exception):
    """Raised when a file's extension isn't supported, or its sniffed
    content doesn't match what its extension claims."""


# Canonical mime_type reported on ParsedDocument, keyed by extension. magic
# alone can't tell .md/.txt/.csv apart (all sniff as text/plain) and docx is
# just a zip, so the extension - not the sniff - decides which parser runs
# and which mime_type is reported; magic only guards against a mismatched
# file being renamed with a supported extension.
_EXTENSION_MIME = {
    ".md": "text/markdown",
    ".txt": "text/plain",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".csv": "text/csv",
    ".json": "application/json",
}

_TEXT_LIKE_EXTENSIONS = {".md", ".txt", ".csv", ".json"}
_DOCX_SNIFFS = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/zip",
}
_MIN_PDF_PAGE_CHARS = 40
_TOP_LEVEL_HEADING_RE = re.compile(r"^#\s+\S.*$", re.MULTILINE)


async def parse_file(path: Path) -> ParsedDocument:
    return await asyncio.to_thread(_parse_file_sync, path)


def _parse_file_sync(path: Path) -> ParsedDocument:
    mime_type = _detect_mime_type(path)
    suffix = path.suffix.lower()

    if suffix in (".md", ".txt"):
        return _parse_text(path, mime_type)
    if suffix == ".pdf":
        return _parse_pdf(path, mime_type)
    if suffix == ".docx":
        return _parse_docx(path, mime_type)
    if suffix == ".csv":
        return _parse_csv(path, mime_type)
    if suffix == ".json":
        return _parse_json(path, mime_type)

    raise UnsupportedFormat(f"Unsupported extension {suffix!r} for {path.name!r}")


def _detect_mime_type(path: Path) -> str:
    suffix = path.suffix.lower()
    expected = _EXTENSION_MIME.get(suffix)
    if expected is None:
        raise UnsupportedFormat(
            f"Unsupported extension {suffix!r} for {path.name!r}; "
            f"supported extensions are {sorted(_EXTENSION_MIME)}"
        )

    sniffed = magic.from_file(str(path), mime=True)
    if not _sniff_matches_extension(suffix, sniffed):
        raise UnsupportedFormat(
            f"{path.name!r} has extension {suffix!r} but its content sniffs as "
            f"{sniffed!r}; refusing to parse"
        )
    return expected


def _sniff_matches_extension(suffix: str, sniffed: str) -> bool:
    # x-empty covers zero-byte files, which magic can't classify by content.
    if sniffed.endswith("x-empty"):
        return True
    if suffix in _TEXT_LIKE_EXTENSIONS:
        return sniffed.startswith("text/") or sniffed == "application/json"
    if suffix == ".pdf":
        return sniffed == "application/pdf"
    if suffix == ".docx":
        return sniffed in _DOCX_SNIFFS
    return False


def _parse_text(path: Path, mime_type: str) -> ParsedDocument:
    text = path.read_text(encoding="utf-8")
    return ParsedDocument(mime_type=mime_type, text=text, segments=_split_by_heading(text))


def _split_by_heading(text: str) -> list[Segment]:
    starts = [m.start() for m in _TOP_LEVEL_HEADING_RE.finditer(text)]
    if not starts:
        return [Segment(char_start=0, char_end=len(text))]

    bounds = starts if starts[0] == 0 else [0, *starts]
    bounds.append(len(text))
    return [Segment(char_start=bounds[i], char_end=bounds[i + 1]) for i in range(len(bounds) - 1)]


def _parse_pdf(path: Path, mime_type: str) -> ParsedDocument:
    reader = PdfReader(str(path))
    parts: list[str] = []
    segments: list[Segment] = []
    cursor = 0
    plumber_doc = None
    try:
        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if len(page_text.strip()) < _MIN_PDF_PAGE_CHARS:
                if plumber_doc is None:
                    plumber_doc = pdfplumber.open(str(path))
                page_text = plumber_doc.pages[page_number - 1].extract_text() or page_text
            parts.append(page_text)
            char_start = cursor
            cursor += len(page_text)
            segments.append(Segment(char_start=char_start, char_end=cursor, page=page_number))
    finally:
        if plumber_doc is not None:
            plumber_doc.close()

    return ParsedDocument(mime_type=mime_type, text="".join(parts), segments=segments)


def _parse_docx(path: Path, mime_type: str) -> ParsedDocument:
    doc = DocxDocument(str(path))
    parts: list[str] = []
    segments: list[Segment] = []
    cursor = 0
    for paragraph in doc.paragraphs:
        para_text = paragraph.text
        if not para_text.strip():
            continue
        if parts:
            parts.append("\n")
            cursor += 1
        char_start = cursor
        parts.append(para_text)
        cursor += len(para_text)
        segments.append(Segment(char_start=char_start, char_end=cursor))

    return ParsedDocument(mime_type=mime_type, text="".join(parts), segments=segments)


def _parse_csv(path: Path, mime_type: str) -> ParsedDocument:
    with path.open(newline="", encoding="utf-8") as f:
        rows = list(csv.reader(f))

    if not rows:
        return ParsedDocument(mime_type=mime_type, text="", segments=[])

    header, data_rows = rows[0], rows[1:]
    parts: list[str] = []
    segments: list[Segment] = []
    cursor = 0
    for row in data_rows:
        row_text = "\n".join(f"{col}: {val}" for col, val in zip(header, row))
        if parts:
            parts.append("\n\n")
            cursor += 2
        char_start = cursor
        parts.append(row_text)
        cursor += len(row_text)
        segments.append(Segment(char_start=char_start, char_end=cursor))

    return ParsedDocument(mime_type=mime_type, text="".join(parts), segments=segments)


def _parse_json(path: Path, mime_type: str) -> ParsedDocument:
    data = json.loads(path.read_text(encoding="utf-8"))

    if isinstance(data, list):
        elements = [(f"$[{i}]", value) for i, value in enumerate(data)]
    elif isinstance(data, dict):
        elements = [(f"$.{key}", value) for key, value in data.items()]
    else:
        elements = [("$", data)]

    parts: list[str] = []
    segments: list[Segment] = []
    cursor = 0
    for json_path, value in elements:
        segment_text = f"{json_path}\n{json.dumps(value, indent=2, ensure_ascii=False)}"
        if parts:
            parts.append("\n\n")
            cursor += 2
        char_start = cursor
        parts.append(segment_text)
        cursor += len(segment_text)
        segments.append(Segment(char_start=char_start, char_end=cursor))

    return ParsedDocument(mime_type=mime_type, text="".join(parts), segments=segments)
