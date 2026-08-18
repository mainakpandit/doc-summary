"""Populates `corpus/demo/` and `corpus/demo2/` with synthetic PM documents
(task_breakdown Step 30 (5)) and, when Postgres is reachable, ingests them
into a `demo` / `demo2` corpus so `make dev`/`make seed` leaves a reviewer
with a corpus "already ingested" (implementation_plan.md section 17) rather
than just files on disk they'd have to upload by hand.

Run with `python backend/scripts/seed_demo.py` (or `make seed`). Writing
files and ingesting them never calls an LLM or embedding provider -- only
`classify`/`extract`/`embed_chunks` do that -- so this needs no API key and
is safe to run in CI or a fresh clone with no `.env` filled in yet. It *does*
need a reachable Postgres for the ingestion half; if `DATABASE_URL` isn't
reachable, the files are still written and a warning is printed instead of
raising, matching the Makefile's `dev` target, which already tolerates this
script failing ("continuing" -- see Makefile).

Idempotent: re-running (every `make dev`) reuses the existing `demo`/`demo2`
Corpus row by name instead of creating a duplicate, and `ingest_file`
already dedupes by content hash within a corpus, so a repeat run is a no-op
past the first.

`corpus/demo/` (6 documents, 4 formats: .md, .csv, .json, .txt) covers three
features across every `doc_type` `agent/nodes/classify.py` recognizes
(prd, techspec, ticket_export, release_notes, meeting_notes, postmortem):

  - "Checkout Redesign": a genuine conflict -- the PRD and the tech spec
    disagree on target_release (2026-Q3 vs 2026-Q4).
  - "Notifications Revamp": an owner-missing case -- mentioned in the ticket
    export and meeting notes, never with an owner, so
    `every_feature_has_owner` (corpus/demo/rules.yaml) should fire on it.
  - "Search Improvements": a clean shipped feature sourced from release
    notes, satisfying `shipped_requires_release_notes`.

`corpus/demo2/` (4 documents, a different format mix: .docx, .pdf, .json,
.txt -- no .md or .csv) covers two different features ("Billing API
Migration", "Search Autocomplete") so a reviewer can rerun the whole flow
against a second corpus with a different declared-format mix, per
WRITEUP.md.
"""

from __future__ import annotations

import asyncio
import json
import sys
from pathlib import Path

from docx import Document as DocxDocument

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

from backend.app.config import get_settings

RULES_YAML_DEMO2 = """\
# Rules playbook for the demo2 corpus. Same predicate vocabulary as
# corpus/demo/rules.yaml -- schema at backend/app/agent/rules/schema.md.

rules:
  - id: every_feature_has_owner
    description: Every feature must have an assigned owner.
    severity: warning
    deterministic:
      op: every_subject_has
      predicate: owner

  - id: every_feature_has_target_release
    description: Every feature must have a target release.
    severity: warning
    deterministic:
      op: every_subject_has
      predicate: target_release
"""


def _pdf_escape(text: str) -> str:
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def _make_pdf_bytes(title: str, lines: list[str]) -> bytes:
    """A minimal, hand-written single-page PDF (no reportlab dependency) --
    same technique already used for backend/tests/fixtures/parsers/sample.pdf,
    just with more lines. Valid enough for pypdf/pdfplumber to extract text
    from, which is all ingestion needs."""
    content_ops = ["BT", "/F1 11 Tf", "72 740 Td", f"({_pdf_escape(title)}) Tj", "0 -24 Td"]
    for i, line in enumerate(lines):
        if i > 0:
            content_ops.append("0 -16 Td")
        content_ops.append(f"({_pdf_escape(line)}) Tj")
    content_ops.append("ET")
    stream = "\n".join(content_ops).encode("latin-1", errors="replace")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 5 0 R >> >> "
            b"/MediaBox [0 0 612 792] /Contents 4 0 R >>"
        ),
        b"<< /Length %d >>\nstream\n" % len(stream) + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode()
        out += body
        out += b"\nendobj\n"

    xref_start = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n" f"startxref\n{xref_start}\n%%EOF"
    ).encode()
    return bytes(out)


def _write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(path)


def _md(text: str) -> bytes:
    return text.encode("utf-8")


def write_demo_corpus(root: Path) -> list[Path]:
    root.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []

    paths.append(root / "prd_checkout_redesign.md")
    paths[-1].write_bytes(
        _md(
            "# Checkout Redesign\n\n"
            "Checkout Redesign is owned by Alice Chen. It targets release "
            "2026-Q3 and is currently in_progress. Open risk: the "
            "third-party payment vendor's sandbox API has been unstable "
            "during integration testing, which could slip the timeline.\n"
        )
    )

    paths.append(root / "techspec_checkout_redesign.md")
    paths[-1].write_bytes(
        _md(
            "# Checkout Redesign -- Technical Design\n\n"
            "This spec covers the backend migration for Checkout Redesign, "
            "owned by Alice Chen. Given the payment vendor sandbox issues "
            "called out by the PRD, engineering now targets release "
            "2026-Q4 instead of the originally proposed date, to leave "
            "time for a vendor bake-off.\n"
        )
    )

    paths.append(root / "tickets_export.csv")
    paths[-1].write_text(
        "feature,status,owner,target_release,summary\n"
        "Checkout Redesign,in_progress,Alice Chen,2026-Q3,"
        "Rebuild the checkout flow on the new payments SDK\n"
        "Notifications Revamp,planned,,2026-Q4,"
        "Consolidate email and push notifications onto one delivery service\n",
        encoding="utf-8",
    )

    paths.append(root / "meeting_notes_planning.txt")
    paths[-1].write_text(
        "Sprint Planning -- Notifications Revamp\n\n"
        "Discussed scope for the Notifications Revamp project. No owner has "
        "been assigned yet pending the reorg; targeting release 2026-Q4. "
        "Open risk raised: the team does not currently have a dedicated "
        "on-call rotation for the new delivery service, which is a gap "
        "before this can ship broadly.\n",
        encoding="utf-8",
    )

    paths.append(root / "release_notes_v25.json")
    paths[-1].write_text(
        json.dumps(
            [
                {
                    "feature": "Search Improvements",
                    "status": "shipped",
                    "owner": "Raj Patel",
                    "released": "2026-Q2",
                    "notes": "Reduced median search latency and improved "
                    "relevance ranking for multi-word queries.",
                }
            ],
            indent=2,
        ),
        encoding="utf-8",
    )

    paths.append(root / "postmortem_search_latency.txt")
    paths[-1].write_text(
        "Postmortem -- Search Improvements Launch\n\n"
        "Owner: Raj Patel. Following the Search Improvements release, a "
        "brief latency regression under peak load was observed and "
        "resolved within two hours by rolling back a caching change. "
        "Open risk: cache warm-up time under cold starts remains "
        "unaddressed.\n",
        encoding="utf-8",
    )

    return paths


def write_demo2_corpus(root: Path) -> list[Path]:
    root.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []

    docx_path = root / "prd_billing_migration.docx"
    _write_docx(
        docx_path,
        "Billing API Migration",
        [
            (
                "Billing API Migration is owned by Jordan Lee. It targets "
                "release 2026-Q4 and is currently planned."
            ),
            (
                "Open risk: the current third-party billing SDK is scheduled "
                "for deprecation before our migration would complete, which "
                "could force an unplanned second migration."
            ),
        ],
    )
    paths.append(docx_path)

    pdf_path = root / "techspec_billing_migration.pdf"
    pdf_path.write_bytes(
        _make_pdf_bytes(
            "Billing API Migration -- Technical Design",
            [
                "Owned by Jordan Lee. Engineering now targets release",
                "2027-Q1 instead of 2026-Q4, to account for the vendor",
                "SDK deprecation timeline called out in the PRD.",
            ],
        )
    )
    paths.append(pdf_path)

    json_path = root / "tickets_search_autocomplete.json"
    json_path.write_text(
        json.dumps(
            [
                {
                    "feature": "Search Autocomplete",
                    "status": "in_progress",
                    "target_release": "2026-Q4",
                    "summary": "Implement typeahead search suggestions "
                    "backed by the existing search index.",
                }
            ],
            indent=2,
        ),
        encoding="utf-8",
    )
    paths.append(json_path)

    txt_path = root / "meeting_notes_q4_planning.txt"
    txt_path.write_text(
        "Q4 Planning Notes\n\n"
        "Billing API Migration: Jordan Lee confirmed as owner, still "
        "targeting a 2027-Q1 release given the SDK deprecation risk.\n\n"
        "Search Autocomplete: no owner assigned yet. Risk raised: "
        "suggestion latency under a cold cache needs a design review "
        "before this can move past in_progress.\n",
        encoding="utf-8",
    )
    paths.append(txt_path)

    return paths


async def _ingest_into_corpus(
    name: str, inbox_path: Path, rules_path: str | None, paths: list[Path]
) -> None:
    from sqlalchemy import select

    from backend.app.db import AsyncSessionLocal
    from backend.app.models.corpus import Corpus
    from backend.app.models.document import Document
    from backend.app.services.ingestion import ingest_file

    async with AsyncSessionLocal() as session:
        corpus = await session.scalar(select(Corpus).where(Corpus.name == name))
        if corpus is None:
            corpus = Corpus(name=name, inbox_path=str(inbox_path), rules_path=rules_path)
            session.add(corpus)
            await session.commit()
            await session.refresh(corpus)
            print(f"  created corpus {name!r} ({corpus.id})")
        else:
            print(f"  reusing existing corpus {name!r} ({corpus.id})")

        for path in paths:
            # Idempotency by filename, not `ingest_file`'s own content-hash
            # dedup: python-docx's zip output embeds a per-save timestamp,
            # so re-running this script regenerates byte-different (but
            # logically identical) .docx files -- content-hash dedup alone
            # would accumulate a near-duplicate Document every run. "Have I
            # already seeded this filename into this corpus" is also just
            # the more natural idempotency question for a seed script to
            # ask (see docs/assumptions.md).
            existing = await session.scalar(
                select(Document).where(
                    Document.corpus_id == corpus.id, Document.filename == path.name
                )
            )
            if existing is not None:
                print(f"    already seeded {path.name} -> document {existing.id}")
                continue

            document = await ingest_file(session, path, corpus.id)
            print(f"    ingested {path.name} -> document {document.id}")


async def _seed_db(settings, demo_paths: list[Path], demo2_paths: list[Path]) -> None:
    inbox_dir = settings.CORPUS_ROOT / "inbox"
    inbox_dir.mkdir(parents=True, exist_ok=True)

    print("Ingesting corpus/demo/ ...")
    await _ingest_into_corpus("demo", inbox_dir, "demo/rules.yaml", demo_paths)

    print("Ingesting corpus/demo2/ ...")
    # A distinct, unwatched inbox path -- services/watcher.py only ever
    # watches settings.CORPUS_ROOT/"inbox" (one physical folder), so demo2
    # gets its own placeholder path rather than colliding with demo's.
    await _ingest_into_corpus(
        "demo2", settings.CORPUS_ROOT / "demo2_inbox", "demo2/rules.yaml", demo2_paths
    )


def main() -> None:
    settings = get_settings()

    demo_root = settings.CORPUS_ROOT / "demo"
    demo2_root = settings.CORPUS_ROOT / "demo2"

    print(f"Writing synthetic demo corpus to {demo_root} ...")
    demo_paths = write_demo_corpus(demo_root)
    for path in demo_paths:
        print(f"  wrote {path.relative_to(REPO_ROOT) if path.is_relative_to(REPO_ROOT) else path}")

    print(f"Writing synthetic demo2 corpus to {demo2_root} ...")
    demo2_paths = write_demo2_corpus(demo2_root)
    for path in demo2_paths:
        print(f"  wrote {path.relative_to(REPO_ROOT) if path.is_relative_to(REPO_ROOT) else path}")

    demo2_rules_path = demo2_root / "rules.yaml"
    if not demo2_rules_path.exists():
        demo2_rules_path.write_text(RULES_YAML_DEMO2, encoding="utf-8")
        print(f"  wrote {demo2_rules_path}")

    try:
        asyncio.run(_seed_db(settings, demo_paths, demo2_paths))
    except Exception as exc:  # pragma: no cover -- exercised manually, needs a live Postgres
        print(f"WARNING: could not ingest demo corpora into the database: {exc}")
        print("Files were still written; run this script again once Postgres is reachable.")


if __name__ == "__main__":
    main()
